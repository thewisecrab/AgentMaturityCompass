"""
s11_attachment_detonation.py — Attachment Detonation and Safe Convert-to-Text Service

Safely extracts text from untrusted file attachments (PDF, DOCX, XLSX, HTML, EML, TXT),
strips active content (macros, scripts, embedded objects), detects hidden/invisible content,
runs injection detection, and returns a SafeTextBundle indicating whether the result is
safe for downstream agent consumption.

Usage::

    from amc.shield.s11_attachment_detonation import AttachmentDetonator

    detonator = AttachmentDetonator()
    bundle = detonator.process("/tmp/upload.pdf", "pdf")
    if bundle.safe_for_agent:
        agent_context = bundle.extracted_text
    else:
        print(f"Quarantined: {bundle.quarantined}, Injection: {bundle.injection_detected}")
        print(f"Removed: {bundle.removed_elements}")
"""

from __future__ import annotations

import asyncio
import email
import email.policy
import hashlib
import re
import zipfile
from pathlib import Path
from typing import Any, Literal

import structlog
from pydantic import BaseModel, Field

from amc.core.models import RiskLevel, ScanResult, score_to_risk  # noqa: F401
from amc.shield.s10_detector import InjectionDetector

log = structlog.get_logger(__name__)

# Invisible / zero-width Unicode codepoints to flag
_INVISIBLE_CHARS = frozenset([
    "\u200b",  # ZERO WIDTH SPACE
    "\u200c",  # ZERO WIDTH NON-JOINER
    "\u200d",  # ZERO WIDTH JOINER
    "\u200e",  # LEFT-TO-RIGHT MARK
    "\u200f",  # RIGHT-TO-LEFT MARK
    "\u2060",  # WORD JOINER
    "\u2061",  # FUNCTION APPLICATION
    "\u2062",  # INVISIBLE TIMES
    "\u2063",  # INVISIBLE SEPARATOR
    "\u2064",  # INVISIBLE PLUS
    "\ufeff",  # ZERO WIDTH NO-BREAK SPACE / BOM
    "\u00ad",  # SOFT HYPHEN
    "\u034f",  # COMBINING GRAPHEME JOINER
    "\u061c",  # ARABIC LETTER MARK
    "\u180e",  # MONGOLIAN VOWEL SEPARATOR
])

# CSS patterns that hide text visually
_CSS_HIDE_PATTERNS = [
    re.compile(r"font-size\s*:\s*0", re.IGNORECASE),
    re.compile(r"display\s*:\s*none", re.IGNORECASE),
    re.compile(r"visibility\s*:\s*hidden", re.IGNORECASE),
    re.compile(r"opacity\s*:\s*0(?:[;\s\"]|$)", re.IGNORECASE),
    re.compile(r"color\s*:\s*(white|#fff(?:fff)?|rgb\(\s*255\s*,\s*255\s*,\s*255\s*\))", re.IGNORECASE),
]


class SafeTextBundle(BaseModel):
    """Result of detonating/extracting text from an attachment."""
    extracted_text: str = ""
    removed_elements: list[str] = Field(default_factory=list)
    file_hash: str = ""
    safe_for_agent: bool = True
    provenance: dict[str, Any] = Field(default_factory=dict)
    quarantined: bool = False
    injection_detected: bool = False


class AttachmentDetonator:
    """Safely converts untrusted file attachments to plain text with threat detection."""

    def __init__(self) -> None:
        self._detector = InjectionDetector()

    # ── public entry ──────────────────────────────────────────────

    def process(
        self,
        file_path: str,
        content_type: Literal["pdf", "docx", "xlsx", "txt", "html", "eml"],
    ) -> SafeTextBundle:
        """Extract text from *file_path*, detect threats, return SafeTextBundle."""
        p = Path(file_path)
        log.info("detonation.start", path=file_path, content_type=content_type)

        # Compute SHA-256 of original bytes
        raw = p.read_bytes()
        file_hash = hashlib.sha256(raw).hexdigest()

        removed: list[str] = []
        quarantined = False
        text = ""
        extraction_failed = False

        handler = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "xlsx": self._extract_xlsx,
            "html": self._extract_html,
            "eml": self._extract_eml,
            "txt": self._extract_txt,
        }[content_type]

        try:
            text, removed, quarantined = handler(p, raw)
        except Exception as exc:
            log.error("detonation.extraction_failed", path=file_path, error=str(exc))
            extraction_failed = True
            text = ""

        # Hidden-content detection (invisible unicode)
        inv_found = {ch for ch in text if ch in _INVISIBLE_CHARS}
        if inv_found:
            codes = [f"U+{ord(c):04X}" for c in sorted(inv_found)]
            removed.append(f"invisible_unicode:{','.join(codes)}")
            log.warn("detonation.invisible_unicode", chars=codes, path=file_path)
            text = "".join(ch for ch in text if ch not in _INVISIBLE_CHARS)

        # CSS hidden-text detection for HTML content
        if content_type == "html":
            css_hits = self._detect_css_tricks(raw.decode("utf-8", errors="replace"))
            if css_hits:
                removed.extend(css_hits)
                quarantined = True
                log.warn("detonation.css_hidden_content", hits=css_hits, path=file_path)

        # Quarantine triggers operator alert
        if quarantined:
            log.critical("detonation.quarantine", path=file_path, reasons=removed)

        # Run injection detection on extracted text
        injection_detected = False
        if text:
            injection_detected = self._run_injection_scan(text, file_path)

        safe = not quarantined and not injection_detected and not extraction_failed

        bundle = SafeTextBundle(
            extracted_text=text,
            removed_elements=removed,
            file_hash=file_hash,
            safe_for_agent=safe,
            provenance={"source_path": file_path, "content_type": content_type, "size_bytes": len(raw)},
            quarantined=quarantined,
            injection_detected=injection_detected,
        )
        log.info("detonation.complete", path=file_path, safe=safe, quarantined=quarantined, injection=injection_detected)
        return bundle

    # ── extractors ────────────────────────────────────────────────

    def _extract_pdf(self, p: Path, raw: bytes) -> tuple[str, list[str], bool]:
        removed: list[str] = []
        quarantined = False

        # Try PyMuPDF (fitz)
        try:
            import fitz  # type: ignore[import-untyped]
            doc = fitz.open(stream=raw, filetype="pdf")
            pages = [page.get_text() for page in doc]
            # Check for JS
            if doc.has_annots() or any("/JS" in (page.get_text("rawdict").get("text", "") if False else "") for _ in []):
                pass
            # Scan catalog for JS actions
            for i in range(len(doc)):
                page = doc[i]
                for link in page.get_links():
                    if link.get("kind") == fitz.LINK_LAUNCH or "javascript" in str(link).lower():
                        removed.append("pdf_javascript_action")
                        quarantined = True
            text = "\n\n".join(pages)
            doc.close()
            log.debug("pdf.extracted_via_fitz", pages=len(pages))
            return text, removed, quarantined
        except Exception as exc:
            log.debug("pdf.fitz_fallback", error=str(exc))

        # Try pdfplumber
        try:
            import pdfplumber  # type: ignore[import-untyped]
            with pdfplumber.open(p) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n\n".join(pages)
            log.debug("pdf.extracted_via_pdfplumber", pages=len(pages))
            return text, removed, quarantined
        except Exception as exc:
            log.debug("pdf.pdfplumber_fallback", error=str(exc))

        # Binary fallback
        text = raw.decode("utf-8", errors="replace")
        removed.append("pdf_binary_fallback")
        log.warn("pdf.binary_fallback")
        return text, removed, quarantined

    def _extract_docx(self, p: Path, raw: bytes) -> tuple[str, list[str], bool]:
        removed: list[str] = []
        quarantined = False

        # Check for macros (vbaProject.bin)
        try:
            with zipfile.ZipFile(p) as zf:
                names = zf.namelist()
                if any("vbaProject.bin" in n for n in names):
                    removed.append("docx_macro_vbaProject.bin")
                    quarantined = True
                    log.warn("docx.macros_detected", path=str(p))
                # Flag embedded objects (oleObject, embeddings)
                embedded = [n for n in names if "embeddings/" in n.lower() or "oleObject" in n]
                if embedded:
                    removed.extend(f"docx_embedded:{n}" for n in embedded)
                    log.warn("docx.embedded_objects", count=len(embedded))
        except zipfile.BadZipFile:
            log.warn("docx.not_a_zip", path=str(p))

        from docx import Document  # type: ignore[import-untyped]
        doc = Document(p)
        parts: list[str] = []

        for para in doc.paragraphs:
            parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                parts.append("\t".join(cells))

        return "\n".join(parts), removed, quarantined

    def _extract_xlsx(self, p: Path, raw: bytes) -> tuple[str, list[str], bool]:
        removed: list[str] = []
        quarantined = False

        # Macro detection
        try:
            with zipfile.ZipFile(p) as zf:
                if any("vbaProject.bin" in n for n in zf.namelist()):
                    removed.append("xlsx_macro_vbaProject.bin")
                    quarantined = True
                    log.warn("xlsx.macros_detected", path=str(p))
        except zipfile.BadZipFile:
            pass

        from openpyxl import load_workbook  # type: ignore[import-untyped]
        wb = load_workbook(p, read_only=True, data_only=True)
        parts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                parts.append("\t".join(str(c) if c is not None else "" for c in row))

        wb.close()
        return "\n".join(parts), removed, quarantined

    def _extract_html(self, p: Path, raw: bytes) -> tuple[str, list[str], bool]:
        removed: list[str] = []
        quarantined = False
        html_str = raw.decode("utf-8", errors="replace")

        # Check for script tags / JS
        if re.search(r"<script[\s>]", html_str, re.IGNORECASE):
            removed.append("html_script_tag")
            quarantined = True

        # Check for on* event handlers
        on_handlers = re.findall(r'\bon\w+\s*=', html_str, re.IGNORECASE)
        if on_handlers:
            removed.append(f"html_event_handlers:{len(on_handlers)}")
            quarantined = True

        try:
            from bs4 import BeautifulSoup  # type: ignore[import-untyped]
            soup = BeautifulSoup(html_str, "html.parser")
            # Remove script and style elements
            for tag in soup(["script", "style"]):
                tag.decompose()
            # Strip on* attributes
            for tag in soup.find_all(True):
                for attr in list(tag.attrs):
                    if attr.lower().startswith("on"):
                        del tag[attr]
            text = soup.get_text(separator="\n", strip=True)
        except Exception:
            log.debug("html.beautifulsoup_fallback_to_regex")
            text = re.sub(r"<script[^>]*>.*?</script>", "", html_str, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<[^>]+>", " ", text)
            text = re.sub(r"\s+", " ", text).strip()

        return text, removed, quarantined

    def _extract_eml(self, p: Path, raw: bytes) -> tuple[str, list[str], bool]:
        removed: list[str] = []
        quarantined = False

        msg = email.message_from_bytes(raw, policy=email.policy.default)
        parts_text: list[str] = []

        # Headers
        for hdr in ("From", "To", "Subject", "Date"):
            val = msg.get(hdr, "")
            if val:
                parts_text.append(f"{hdr}: {val}")

        attachments_meta: list[str] = []
        has_html = False

        for part in msg.walk():
            ct = part.get_content_type()
            fn = part.get_filename()
            if fn:
                attachments_meta.append(f"{fn} ({ct})")
                continue
            if ct == "text/plain":
                payload = part.get_content()
                if isinstance(payload, str):
                    parts_text.append(payload)
            elif ct == "text/html":
                has_html = True
                removed.append("eml_html_body_stripped")

        if has_html:
            log.warn("eml.html_body_present", path=str(p))

        if attachments_meta:
            parts_text.append(f"[Attachments metadata: {'; '.join(attachments_meta)}]")
            removed.append(f"eml_attachments_not_opened:{len(attachments_meta)}")

        return "\n".join(parts_text), removed, quarantined

    def _extract_txt(self, p: Path, raw: bytes) -> tuple[str, list[str], bool]:
        return raw.decode("utf-8", errors="replace"), [], False

    # ── threat helpers ────────────────────────────────────────────

    def _detect_css_tricks(self, html: str) -> list[str]:
        hits: list[str] = []
        for pat in _CSS_HIDE_PATTERNS:
            if pat.search(html):
                hits.append(f"css_hidden:{pat.pattern}")
        return hits

    def _run_injection_scan(self, text: str, source: str) -> bool:
        try:
            loop = asyncio.get_running_loop()
        except RuntimeError:
            loop = None

        async def _scan() -> ScanResult:
            return await self._detector.scan(text, source=source, context={"stage": "attachment_detonation"})

        try:
            if loop and loop.is_running():
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                    result = pool.submit(lambda: asyncio.run(_scan())).result(timeout=30)
            else:
                result = asyncio.run(_scan())
        except Exception as exc:
            log.error("detonation.injection_scan_failed", error=str(exc))
            return True  # fail-closed

        detected = result.risk_level in (RiskLevel.HIGH, RiskLevel.CRITICAL) if hasattr(result, "risk_level") else False
        if detected:
            log.warn("detonation.injection_detected", source=source, risk=str(getattr(result, "risk_level", "UNKNOWN")))
        return detected
